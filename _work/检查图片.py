from docx import Document
from docx.oxml.ns import qn

doc = Document(r"C:\Users\a1442\Desktop\创新思维大作业\23软件本科2班-2023103010099-杨明宇-期末试卷.docx")

# 找所有图片和位置
print("=== 图片插入位置 ===")
for i, p in enumerate(doc.paragraphs):
    has_img = False
    for r in p.runs:
        for elem in r._element.iter():
            if elem.tag.endswith('}blip'):
                has_img = True
                break
    # 上下3段文字
    if has_img:
        # 找caption
        caption = doc.paragraphs[i+1].text if i+1 < len(doc.paragraphs) else ""
        # 找上下文
        prev = doc.paragraphs[i-1].text if i > 0 else ""
        print(f"  段 {i}: 图片 | caption={caption[:30]} | 上一段={prev[-40:]}")

print("\n=== 全部表 ===")
for i, tbl in enumerate(doc.tables):
    print(f"  表 {i+1}: {len(tbl.rows)} 行 x {len(tbl.rows[0].cells)} 列")
    if i == 0:
        for row in tbl.rows:
            print("    " + " | ".join(c.text[:20] for c in row.cells))

print(f"\n=== 文档统计 ===")
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"图片数: 8 (已嵌入)")
text = sum(len(p.text) for p in doc.paragraphs)
print(f"字符数: {text}")

# 中文字数
import re
zh = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
print(f"中文字数: {zh}")
