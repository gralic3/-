from docx import Document
doc = Document(r"C:\Users\a1442\Desktop\创新思维大作业\23软件本科2班-2023103010099-杨明宇-期末试卷.docx")
print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")
print()
print("=== 前 30 段 ===")
for i, p in enumerate(doc.paragraphs[:30]):
    txt = p.text.strip()
    if txt:
        print(f"[{i:3d}] {txt[:80]}")

# 检查图片
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        for elem in r._element.iter():
            if elem.tag.endswith('}blip'):
                img_count += 1
print(f"\n实际嵌入图片数: {img_count}")

print("\n=== 表格 ===")
for i, tbl in enumerate(doc.tables):
    print(f"表 {i+1}: {len(tbl.rows)}x{len(tbl.rows[0].cells) if tbl.rows else 0}")
    for row in tbl.rows[:3]:
        print("  | " + " | ".join(c.text for c in row.cells) + " |")
    if len(tbl.rows) > 3:
        print(f"  ... 共 {len(tbl.rows)} 行")

# 统计纯文字 (排除表格和图)
text_chars = sum(len(p.text) for p in doc.paragraphs)
print(f"\n纯文字字符数: {text_chars}")
