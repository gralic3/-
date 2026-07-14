"""
把论文初稿 + 图表按章节结构写入 docx 答题纸
输入: 论文初稿.md + figures/
输出: 23软件本科2班-2023103010099-杨明宇-期末试卷.docx (已填好内容)
"""
import os
import re
import shutil
import markdown
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:\Users\a1442\Desktop\创新思维大作业"
WORK = os.path.join(ROOT, "_work")
FIG = os.path.join(WORK, "figures")
SRC_DOCX = os.path.join(ROOT, "25-26-02-23软件本科《创新思维培养与创业管理》期末A卷-大作业-答题纸.docx")
DST_DOCX = os.path.join(ROOT, "23软件本科2班-2023103010099-杨明宇-期末试卷.docx")
MD_FILE = os.path.join(WORK, "论文初稿.md")

# 1. 复制原 docx 作为模板
# ============================================================
print("[1/5] 复制原 docx 模板 ...")
import os
if os.path.exists(DST_DOCX):
    os.remove(DST_DOCX)
shutil.copy2(SRC_DOCX, DST_DOCX)
doc = Document(DST_DOCX)

# 提取封面信息 (从模板里读取)
cover_info = {
    "课程名": "《创新思维培养与创业管理》",
    "姓名": "杨明宇",
    "学号": "2023103010099",
    "专业": "软件工程",
    "班级": "23软件本科2班",
    "日期": "2026年6月",
}

# ============================================================
# 1.5. 清空 docx 的所有表格 (模板里的题目/姓名/学号/样表)
# ============================================================
print("[1.5/5] 清空 docx 原表格 ...")
# python-docx 的 doc.tables 删除比较复杂, 用底层 XML
body = doc.element.body
for tbl in list(body.findall(qn('w:tbl'))):
    body.remove(tbl)
print("    已清空所有原表格")

# ============================================================
# 2. 解析 markdown 论文
# ============================================================
print("[2/5] 解析 markdown ...")
with open(MD_FILE, "r", encoding="utf-8") as f:
    md_text = f.read()
# 找到 "## 摘要" 所在行
lines = md_text.split("\n")
start_idx = next(i for i, l in enumerate(lines) if l.strip().startswith("## 摘要"))
md_body = "\n".join(lines[start_idx:])

# 提取标题（第一个 #）
title_line = next(l for l in lines if l.startswith("# "))
title = title_line.lstrip("# ").strip()
print(f"    论文标题: {title}")

# 把 markdown 转成 HTML
html = markdown.markdown(md_body, extensions=["extra", "codehilite", "tables"])
# 解析 HTML（用简单正则）
import re
from html import unescape

# 提取所有块级元素
def parse_blocks(html):
    """解析 HTML 为 [(type, content), ...] 块列表"""
    blocks = []
    # 先处理图片
    img_pattern = r'<img[^>]*alt="([^"]*)"[^>]*src="([^"]*)"[^>]*/?>'
    # 简化: 找到 <img> 后替换为 [IMG: alt | src]
    # 然后按段落切分
    # 先把表格抽出来
    # 这里用 BeautifulSoup 替代
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'pre']):
        if elem.name.startswith('h'):
            level = int(elem.name[1])
            text = elem.get_text().strip()
            blocks.append(('h', level, text))
        elif elem.name == 'p':
            text = elem.get_text().strip()
            # 检查是否包含公式 (用 $...$ 或 $$...$$)
            if '$$' in elem.get_text():
                # 公式特殊处理
                blocks.append(('formula', elem.get_text().strip()))
            else:
                blocks.append(('p', text))
        elif elem.name in ('ul', 'ol'):
            items = [li.get_text().strip() for li in elem.find_all('li')]
            blocks.append(('list', items, elem.name == 'ol'))
        elif elem.name == 'table':
            rows = []
            for tr in elem.find_all('tr'):
                cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                rows.append(cells)
            blocks.append(('table', rows))
    return blocks

try:
    blocks = parse_blocks(html)
except ImportError:
    print("    [警告] bs4 未安装, 用简单正则解析")
    blocks = []

# ============================================================
# 3. 清空 docx 正文 (保留封面表格 1, 删除其他所有)
# ============================================================
print("[3/5] 清空 docx 现有正文 ...")
# 删除所有段落（包括封面之后的）
# 保留: 文档最前面的 0-2 个段（标题"《创新思维培养与创业管理》"）
# 直接全部删除
all_paragraphs = list(doc.paragraphs)
for p in all_paragraphs:
    p._element.getparent().remove(p._element)
# 删除所有表格（保留封面表格除外? 不, 全部删了重新建）
# 实际策略: 保留第一个表格 (封面) 不动, 删其他
# 简化: 全部删掉, 我们用纯段落+表重建
# 但要保留封面信息 (题目、姓名、学号、专业、班级)
# 这里采用"重建"策略: 清空所有段落, 重新插入
# 但封面信息已经在 doc 里了, 让我们从原 docx 提取
print("    已清空所有段落")

# ============================================================
# 3.5. 重建封面
# ============================================================
print("[3.5/5] 重建封面 ...")

def add_centered_text(doc, text, size=14, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "宋体")
    rPr.append(rFonts)
    return p

def add_field_table(doc, label, value):
    """添加一个简单的标签-值对"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Cm(3)
    tbl.columns[1].width = Cm(10)
    cell0 = tbl.rows[0].cells[0]
    cell0.text = label
    cell0.width = Cm(3)
    cell1 = tbl.rows[0].cells[1]
    cell1.text = value
    cell1.width = Cm(10)
    return tbl

# 封面: 居中标题
add_centered_text(doc, "", 12)
add_centered_text(doc, "广 东 科 技 学 院", size=22, bold=True)
add_centered_text(doc, "Guangdong University of Science and Technology", size=12)
add_centered_text(doc, "", 12)
add_centered_text(doc, "《创新思维培养与创业管理》", size=16, bold=True)
add_centered_text(doc, "期末考试课程论文", size=14)
add_centered_text(doc, "", 12)
# 题目 (大号, 居中, 加粗)
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("题  目：" + title)
run.font.size = Pt(16)
run.bold = True
run.font.name = "宋体"
rPr = run._element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), "宋体")
rPr.append(rFonts)
add_centered_text(doc, "", 12)

# 信息表
info_table = doc.add_table(rows=6, cols=2)
info_table.style = "Table Grid"
info_table.autofit = False
info_data = [
    ("姓    名", cover_info["姓名"]),
    ("学    号", cover_info["学号"]),
    ("专    业", cover_info["专业"]),
    ("班    级", cover_info["班级"]),
    ("学    院", "计算机学院"),
    ("日    期", cover_info["日期"]),
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    # 加粗左侧
    for run in info_table.rows[i].cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(12)
    for run in info_table.rows[i].cells[1].paragraphs[0].runs:
        run.font.size = Pt(12)

# 分页
doc.add_page_break()
print("    封面已重建")

# ============================================================
# 4. 写入新内容
# ============================================================
print(f"[4/5] 写入 {len(blocks)} 个块 ...")

def set_para_style(p, level=None, bold=False, size=12, align=None):
    """设置段落格式"""
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        run.font.size = Pt(size)
        run.bold = bold

def add_picture(doc, img_path, width_cm=12, caption=None):
    """添加图片 + 标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        run = p.add_run(f"[图片未找到: {img_path}]")
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(10)
            r.bold = True
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

# 图表与插入位置的映射 (从论文里读)
IMG_MAP = {
    "图 3.1": os.path.join(FIG, "3-1-arch.png"),
    "图 3.2": os.path.join(FIG, "3-2-usecase.png"),
    "图 4.1": os.path.join(FIG, "4-1-dataflow.png"),
    "图 4.2": os.path.join(FIG, "4-2-kg.png"),
    "图 4.3": os.path.join(FIG, "4-3-tsne.png"),
    "图 5.1": os.path.join(FIG, "5-1-dashboard.png"),
    "图 5.2": os.path.join(FIG, "5-2-roc.png"),
    "图 5.3": os.path.join(FIG, "5-3-shap-summary.png"),
    "图 5.4": os.path.join(FIG, "5-4-shap-force.png"),
}

added_imgs = set()
# 收集每张图第一次出现在哪个 block
# 一个 block 可能提到多张图, 都要记录
img_first_block = {}  # img_key -> bi
for bi, block in enumerate(blocks):
    if block[0] != 'p':
        continue
    text = block[1]
    # 找该 block 里所有图引用
    figs_in_block = list(set(re.findall(r'图\s*\d+\.\d+', text)))
    for k in figs_in_block:
        if k in IMG_MAP and k not in img_first_block:
            img_first_block[k] = bi

for bi, block in enumerate(blocks):
    if block[0] == 'h':
        level = block[1]
        text = block[2]
        if level == 2:  # ## 摘要, ## 1 引言, ## 参考文献
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif level == 3:  # ### 3.1 需求分析
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        elif level == 4:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    elif block[0] == 'p':
        text = block[1]
        if not text:
            continue
        # 检查是否需要先插入图片 (按 block 索引)
        # 当前 block 第一次出现的所有未插入的图
        figs_to_insert_now = []
        for k, v in img_first_block.items():
            if v == bi and k not in added_imgs:
                figs_to_insert_now.append(k)
        for img_key in figs_to_insert_now:
            img_path = IMG_MAP[img_key]
            if os.path.exists(img_path):
                add_picture(doc, img_path, width_cm=12, caption=img_key + " 系统相关示意图")
            else:
                # 架构图占位 (后续用 draw.io 画)
                placeholder = doc.add_paragraph()
                placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = placeholder.add_run(f"[ {img_key} 待 draw.io 绘制 — 文件: {os.path.basename(img_path)} ]")
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                # 占位框
                cp = doc.add_paragraph(img_key + "  (待绘制)")
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cp.runs:
                    r.font.size = Pt(10)
                    r.bold = True
                    r.italic = True
            added_imgs.add(img_key)
        # 检查是否提到表 X.Y
        if "表 5.1" in text and "table_5_1" not in added_imgs:
            # 插表 5.1
            tbl = doc.add_table(rows=7, cols=6)
            tbl.style = "Table Grid"
            header = ["模型", "AUROC", "F1@10%", "Precision@K", "Recall@K", "提前预警天数(天)"]
            for i, h in enumerate(header):
                tbl.rows[0].cells[i].text = h
            data = [
                ["LR", "0.65", "0.29", "0.37", "0.24", "—"],
                ["XGBoost (无 KG)", "0.59", "0.16", "0.20", "0.13", "14"],
                ["XGBoost + 行为特征", "0.64", "0.28", "0.36", "0.23", "21"],
                ["本文 (XGBoost + 行为 + KG)", "0.87", "0.61", "0.78", "0.51", "28"],
            ]
            for r_idx, row in enumerate(data, 1):
                for c_idx, cell in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell
            # 加表标题
            cp = doc.add_paragraph("表 5.1  不同模型性能对比")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(10)
                r.bold = True
            added_imgs.add("table_5_1")
        if "表 5.2" in text and "table_5_2" not in added_imgs:
            tbl = doc.add_table(rows=6, cols=3)
            tbl.style = "Table Grid"
            header = ["模型变体", "AUROC", "较完整模型下降"]
            for i, h in enumerate(header):
                tbl.rows[0].cells[i].text = h
            data = [
                ["完整模型 (XGBoost+行为+KG)", "0.87", "—"],
                ["- 移除 KG 特征", "0.59", "-0.28"],
                ["- 注意力池化 -> 均值池化", "0.64", "-0.23"],
                ["- 7 天窗口 -> 3 天窗口", "0.85", "-0.02"],
                ["- XGBoost -> 随机森林", "0.80", "-0.07"],
            ]
            for r_idx, row in enumerate(data, 1):
                for c_idx, cell in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell
            cp = doc.add_paragraph("表 5.2  消融实验结果")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(10)
                r.bold = True
            added_imgs.add("table_5_2")
        # 写段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        # 把 **加粗** 转为普通文本 (这里简化处理)
        text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        run = p.add_run(text_clean)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    elif block[0] == 'formula':
        text = block[1]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.italic = True
    elif block[0] == 'list':
        items = block[1]
        ordered = block[2]
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph()
            prefix = f"{idx}. " if ordered else "• "
            run = p.add_run(prefix + item)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    elif block[0] == 'table':
        rows = block[1]
        if not rows:
            continue
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                if c_idx < len(tbl.rows[r_idx].cells):
                    tbl.rows[r_idx].cells[c_idx].text = cell

# 插入仍未出现的图片（兜底）
for img_key, img_path in IMG_MAP.items():
    if img_key not in added_imgs:
        add_picture(doc, img_path, width_cm=12, caption=f"{img_key} 系统相关示意图")
        added_imgs.add(img_key)

# ============================================================
# 5. 保存
# ============================================================
print("[5/5] 保存 docx ...")
doc.save(DST_DOCX)
print(f"\n[完成] 已写入: {DST_DOCX}")
print(f"   文档大小: {os.path.getsize(DST_DOCX) // 1024} KB")
print(f"   已插入图片: {len(added_imgs)} 张")

# 字数统计
total_chars = sum(len(p.text) for p in doc.paragraphs)
print(f"   总字数: {total_chars} 字")
