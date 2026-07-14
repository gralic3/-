from docx import Document
doc = Document(r"C:\Users\a1442\Desktop\创新思维大作业\25-26-02-23软件本科《创新思维培养与创业管理》期末A卷-大作业-答题纸.docx")
total_chars = 0
total_zh = 0
for p in doc.paragraphs:
    t = p.text
    total_chars += len(t)
    for c in t:
        if '\u4e00' <= c <= '\u9fff':
            total_zh += 1

# 加上表格里的字
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            t = cell.text
            total_chars += len(t)
            for c in t:
                if '\u4e00' <= c <= '\u9fff':
                    total_zh += 1

# 图片
img = 0
for p in doc.paragraphs:
    for r in p.runs:
        for elem in r._element.iter():
            if elem.tag.endswith('}blip'):
                img += 1

print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"图片数: {img} ← 目标: 0")
print(f"总字符数: {total_chars}")
print(f"中文字数: {total_zh}")
print(f"估计正文字数: {total_zh} 字 (在 2500-4000 区间)")

# 看前几段
print("\n=== 前 15 段 ===")
for i, p in enumerate(doc.paragraphs[:15]):
    txt = p.text.strip()
    if txt:
        print(f"[{i:2d}] {txt[:70]}")

print("\n=== 3 张表 ===")
for i, tbl in enumerate(doc.tables):
    print(f"表 {i+1}: {len(tbl.rows)} 行 x {len(tbl.rows[0].cells)} 列")
    if i > 0:  # 详细展示 2 张数据表
        for row in tbl.rows[:3]:
            print("  | " + " | ".join(c.text[:15] for c in row.cells) + " |")
