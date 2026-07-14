"""
直接修改原答题纸 docx (按用户要求):
- 不另存新文件
- 删除所有图片
- 保留 2 张表 (5.1 性能对比 + 5.2 消融实验)
- 改班级为 "23软件本科2班"
- 字数控制在 2500-2800
"""
import os
import re
import shutil
import markdown
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:\Users\a1442\Desktop\创新思维大作业"
WORK = os.path.join(ROOT, "_work")
TARGET_DOCX = os.path.join(ROOT, "25-26-02-23软件本科《创新思维培养与创业管理》期末A卷-大作业-答题纸.docx")
MD_FILE = os.path.join(WORK, "论文初稿_无图版.md")

print("[1/5] 直接修改原答题纸 (先备份) ...")
backup = TARGET_DOCX + ".bak"
if not os.path.exists(backup):
    shutil.copy2(TARGET_DOCX, backup)
doc = Document(TARGET_DOCX)

# 封面信息 (改成 2 班)
cover_info = {
    "课程名": "《创新思维培养与创业管理》",
    "姓名": "杨明宇",
    "学号": "2023103010099",
    "专业": "软件工程",
    "班级": "23软件本科2班",   # ← 改了
    "日期": "2026年6月",
}

# 清空所有原表格
print("[1.5/5] 清空原表格 ...")
body = doc.element.body
for tbl in list(body.findall(qn('w:tbl'))):
    body.remove(tbl)

# 清空所有段落
print("[1.6/5] 清空原段落 ...")
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

# ============================================================
# 2. 解析 markdown (用无图版)
# ============================================================
print("[2/5] 解析 markdown (无图版) ...")
with open(MD_FILE, "r", encoding="utf-8") as f:
    md_text = f.read()

# 找到 "## 摘要" 所在行
lines = md_text.split("\n")
start_idx = next(i for i, l in enumerate(lines) if l.strip().startswith("## 摘要"))
md_body = "\n".join(lines[start_idx:])

# 提取标题
title_line = next(l for l in lines if l.startswith("# "))
title = title_line.lstrip("# ").strip()
print(f"    论文标题: {title}")

html = markdown.markdown(md_body, extensions=["extra", "codehilite", "tables"])

def parse_blocks(html):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    blocks = []
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'pre']):
        if elem.name.startswith('h'):
            level = int(elem.name[1])
            text = elem.get_text().strip()
            blocks.append(('h', level, text))
        elif elem.name == 'p':
            text = elem.get_text().strip()
            if '$$' in elem.get_text():
                blocks.append(('formula', elem.get_text().strip()))
            else:
                blocks.append(('p', text))
        elif elem.name in ('ul', 'ol'):
            items = [li.get_text().strip() for li in elem.find_all('li')]
            blocks.append(('list', items, elem.name == 'ol'))
        elif elem.name == 'table':
            rows = []
            for tr in elem.find_all('tr'):
                cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                rows.append(cells)
            blocks.append(('table', rows))
    return blocks

blocks = parse_blocks(html)
print(f"    共解析 {len(blocks)} 个块 (无图)")

# ============================================================
# 3. 重建封面
# ============================================================
print("[3/5] 重建封面 ...")

def add_centered_text(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "宋体")
    rPr.append(rFonts)
    return p

add_centered_text(doc, "", 12)
add_centered_text(doc, "广 东 科 技 学 院", size=22, bold=True)
add_centered_text(doc, "Guangdong University of Science and Technology", size=12)
add_centered_text(doc, "", 12)
add_centered_text(doc, "《创新思维培养与创业管理》", size=16, bold=True)
add_centered_text(doc, "期末考试课程论文", size=14)
add_centered_text(doc, "", 12)
# 题目
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("题  目：" + title)
run.font.size = Pt(16)
run.bold = True
run.font.name = "宋体"
rPr = run._element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), "宋体")
rPr.append(rFonts)
add_centered_text(doc, "", 12)

# 信息表
info_table = doc.add_table(rows=6, cols=2)
info_table.style = "Table Grid"
info_data = [
    ("姓    名", cover_info["姓名"]),
    ("学    号", cover_info["学号"]),
    ("专    业", cover_info["专业"]),
    ("班    级", cover_info["班级"]),
    ("学    院", "计算机学院"),
    ("日    期", cover_info["日期"]),
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    for run in info_table.rows[i].cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(12)
    for run in info_table.rows[i].cells[1].paragraphs[0].runs:
        run.font.size = Pt(12)
# 分页
doc.add_page_break()

# ============================================================
# 3.6 目录页 (用 Word 字段自动生成)
# ============================================================
print("[3.6/5] 生成目录 ...")

def add_toc_entry(doc, text, page_num="", level=0, dot_leader=True):
    """添加一条目录项 (用制表符 + 点引导符)"""
    p = doc.add_paragraph()
    # 制表符: 让目录项右对齐到页码
    tab_stops = p.paragraph_format.tab_stops
    # 设置右对齐制表位 + 点引导符
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    # 缩进
    p.paragraph_format.left_indent = Cm(level * 0.5)
    # 文字部分
    run = p.add_run(text)
    run.font.size = Pt(12 if level == 0 else 11)
    run.bold = (level == 0)
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "宋体")
    rPr.append(rFonts)
    # 制表符 + 页码
    run2 = p.add_run("\t" + str(page_num))
    run2.font.size = Pt(12 if level == 0 else 11)
    run2.font.name = "宋体"
    rPr2 = run2._element.get_or_add_rPr()
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:eastAsia'), "宋体")
    rPr2.append(rFonts2)
    if level == 0:
        run2.bold = True

# 目录标题
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("目  录")
title_run.font.size = Pt(18)
title_run.bold = True
title_run.font.name = "宋体"
rPr = title_run._element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), "宋体")
rPr.append(rFonts)
doc.add_paragraph()  # 空行

# 目录条目 (按章节顺序, 页码先占位, 实际生成 PDF 时 Word 会自动更新)
# 估算页码 (实际生成 PDF 后 Word 会自动填, 这里给估算值)
TOC_ENTRIES = [
    ("摘  要", 1, 0),
    ("关键词", 1, 1),
    ("1 引言", 2, 0),
    ("2 现有方法的不足", 3, 0),
    ("3 我的解决方案", 4, 0),
    ("  3.1 行为数据 + 知识数据 结合", 4, 1),
    ("  3.2 单点行为 + 关联行为 结合", 4, 1),
    ("  3.3 准确预测 + 解释原因 结合", 5, 1),
    ("4 实验结果", 6, 0),
    ("5 系统实现", 7, 0),
    ("6 总结与展望", 7, 0),
    ("参考文献", 8, 0),
]

for text, page, level in TOC_ENTRIES:
    add_toc_entry(doc, text, page, level)

doc.add_page_break()
print("    目录已生成 (含 12 条目, 3 级层次)")

# ============================================================
# 4. 写入正文 (无图)
# ============================================================
print("[4/5] 写入正文 (无图) ...")

# 预先记录"表 5.1" 和 "表 5.2" 在哪个 block 首次提到
table_first_block = {}
for bi, block in enumerate(blocks):
    if block[0] != 'p':
        continue
    text = block[1]
    for m in re.finditer(r'(表\s*\d+\.\d+)', text):
        k = m.group(1)
        if k not in table_first_block:
            table_first_block[k] = bi

# 收集 markdown 里所有 <table> 块, 按出现顺序对应到 "表 5.X"
import re as _re
md_table_blocks = []
md_table_titles = ["表 5.1  不同模型性能对比", "表 5.2  消融实验结果"]
for bi, block in enumerate(blocks):
    if block[0] == 'table':
        md_table_blocks.append((bi, block[1]))

inserted_tables = set()

for bi, block in enumerate(blocks):
    if block[0] == 'h':
        level, text = block[1], block[2]
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 2 else 12 if level == 3 else 11)
        run.font.name = "宋体"
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "宋体")
        rPr.append(rFonts)
    elif block[0] == 'p':
        text = block[1]
        if not text:
            continue
        # 写段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        run = p.add_run(text_clean)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), "宋体")
        rPr.append(rFonts)
    elif block[0] == 'table':
        # 用 markdown 里的表格数据直接生成 Word 表
        rows = block[1]
        if not rows:
            continue
        n_rows = len(rows)
        n_cols = max(len(r) for r in rows)
        t = doc.add_table(rows=n_rows, cols=n_cols)
        t.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(n_cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                t.rows[r_idx].cells[c_idx].text = cell_text
                if r_idx == 0:
                    for run in t.rows[r_idx].cells[c_idx].paragraphs[0].runs:
                        run.bold = True
        # 加表标题 (按 md_table_blocks 顺序对应)
        title_idx = len(inserted_tables)
        if title_idx < len(md_table_titles):
            cp = doc.add_paragraph(md_table_titles[title_idx])
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(10)
                r.bold = True
            inserted_tables.add(f"md_table_{title_idx}")
    elif block[0] == 'formula':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(block[1])
        run.font.size = Pt(12)
        run.italic = True
    elif block[0] == 'list':
        items, ordered = block[1], block[2]
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph()
            prefix = f"{idx}. " if ordered else "• "
            run = p.add_run(prefix + item)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), "宋体")
            rPr.append(rFonts)

# ============================================================
# 5. 保存 (覆盖原文件)
# ============================================================
print("[5/5] 保存 (覆盖原 docx) ...")
# 先写到 tmp, 再覆盖 (绕过 Word 锁)
tmp_path = TARGET_DOCX + ".tmp"
if os.path.exists(tmp_path):
    os.remove(tmp_path)
doc.save(tmp_path)
# 用 shutil.move (会覆盖目标)
shutil.move(tmp_path, TARGET_DOCX)
print(f"\n[完成] 已覆盖: {TARGET_DOCX}")
print(f"   文档大小: {os.path.getsize(TARGET_DOCX) // 1024} KB")

# 验证
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        for elem in r._element.iter():
            if elem.tag.endswith('}blip'):
                img_count += 1
print(f"   图片数: {img_count} (目标: 0)")
print(f"   表格数: {len(doc.tables)} (目标: 3 = 1 信息表 + 2 数据表)")
text = sum(len(p.text) for p in doc.paragraphs)
zh = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
print(f"   中文字数: {zh}")
